"""Apply Word-native table geometry and visual styling to a DOCX report.

Pandoc already converts Markdown pipe tables to Word tables, but the generated
OOXML can rely on percentage widths and default table styling. This postprocess
adds fixed table width, column grid widths, borders, cell margins, and header
shading so tables remain readable in Word and Google Docs exports.
"""

from __future__ import annotations

import sys
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


TOTAL_TABLE_DXA = 9029  # A4 width 11909 dxa minus 1 inch margins on both sides.
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"


def _get_or_add(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _ensure_font_content_types(input_path: Path) -> Path:
    """Return a DOCX path whose Content_Types includes embedded font defaults."""
    with zipfile.ZipFile(input_path) as zin:
        names = zin.namelist()
        needs_ttf = any(name.lower().endswith(".ttf") for name in names)
        if not needs_ttf:
            return input_path
        content_types = zin.read("[Content_Types].xml")

    ET.register_namespace("", CT_NS)
    root = ET.fromstring(content_types)
    has_ttf = any(
        child.tag == f"{{{CT_NS}}}Default" and child.attrib.get("Extension", "").lower() == "ttf"
        for child in root
    )
    if has_ttf:
        return input_path

    default = ET.Element(f"{{{CT_NS}}}Default")
    default.set("Extension", "ttf")
    default.set("ContentType", "application/x-font-ttf")
    root.append(default)
    fixed_content_types = ET.tostring(root, encoding="utf-8", xml_declaration=True)

    tmp = Path(tempfile.mkstemp(suffix=".docx")[1])
    with zipfile.ZipFile(input_path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = fixed_content_types
            zout.writestr(item, data)
    return tmp


def _set_table_width(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = _get_or_add(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = _get_or_add(tbl_pr, "w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "0")

    tbl_layout = _get_or_add(tbl_pr, "w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")

    for old_grid in tbl.findall(qn("w:tblGrid")):
        tbl.remove(old_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    tbl.insert(1, tbl_grid)


def _set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = _get_or_add(tbl_pr, "w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = _get_or_add(borders, f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "D9E2EC")


def _set_table_cell_margins(table, margin: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = _get_or_add(tbl_pr, "w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        elem = _get_or_add(margins, f"w:{side}")
        elem.set(qn("w:w"), str(margin))
        elem.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = _get_or_add(tc_pr, "w:tcW")
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width))
    v_align = _get_or_add(tc_pr, "w:vAlign")
    v_align.set(qn("w:val"), "center")


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = _get_or_add(tc_pr, "w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def _has_drawing(cell) -> bool:
    return bool(cell._tc.xpath(".//w:drawing"))


def _plain_text(cell) -> str:
    return " ".join(cell.text.split())


def _column_widths(table) -> list[int]:
    ncols = max(len(row.cells) for row in table.rows)
    if ncols <= 1:
        return [TOTAL_TABLE_DXA]

    has_image_col = [
        any(idx < len(row.cells) and _has_drawing(row.cells[idx]) for row in table.rows)
        for idx in range(ncols)
    ]

    if ncols == 2 and has_image_col[0] and not has_image_col[1]:
        return [3600, TOTAL_TABLE_DXA - 3600]

    scores: list[float] = []
    for idx in range(ncols):
        max_len = 1
        for row in table.rows:
            if idx < len(row.cells):
                max_len = max(max_len, len(_plain_text(row.cells[idx])))
        score = min(max_len, 55)
        if has_image_col[idx]:
            score = max(score, 28)
        scores.append(float(score))

    min_width = 900 if ncols >= 5 else 1200
    available = TOTAL_TABLE_DXA - min_width * ncols
    if available < 0:
        base = TOTAL_TABLE_DXA // ncols
        widths = [base] * ncols
        widths[-1] += TOTAL_TABLE_DXA - sum(widths)
        return widths

    score_total = sum(scores) or 1.0
    widths = [min_width + round(available * score / score_total) for score in scores]
    widths[-1] += TOTAL_TABLE_DXA - sum(widths)
    return widths


def _format_cell_text(cell, is_header: bool, ncols: int) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        if is_header:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif len(_plain_text(cell)) <= 16 and not _has_drawing(cell):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for run in paragraph.runs:
            run.font.size = Pt(8 if ncols >= 5 else 8.5)
            if is_header:
                run.bold = True


def style_docx_tables(input_path: Path, output_path: Path) -> None:
    readable_input = _ensure_font_content_types(input_path)
    doc = Document(readable_input)

    for shape in doc.inline_shapes:
        max_width = Inches(6.2)
        if shape.width and shape.width > max_width:
            ratio = max_width / shape.width
            shape.width = max_width
            shape.height = int(shape.height * ratio)

    for table in doc.tables:
        if not table.rows:
            continue

        widths = _column_widths(table)
        _set_table_width(table, widths)
        _set_table_borders(table)
        _set_table_cell_margins(table)

        try:
            table.style = "Table Grid"
        except KeyError:
            pass

        for row_idx, row in enumerate(table.rows):
            is_header = row_idx == 0
            if is_header:
                tr_pr = row._tr.get_or_add_trPr()
                tbl_header = _get_or_add(tr_pr, "w:tblHeader")
                tbl_header.set(qn("w:val"), "true")
            for cell_idx, cell in enumerate(row.cells):
                width = widths[min(cell_idx, len(widths) - 1)]
                _set_cell_width(cell, width)
                if is_header:
                    _shade_cell(cell, "EAF2F8")
                _format_cell_text(cell, is_header, len(widths))

    doc.save(output_path)


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("usage: style_docx_tables.py input.docx output.docx")
    style_docx_tables(Path(sys.argv[1]), Path(sys.argv[2]))


if __name__ == "__main__":
    main()
